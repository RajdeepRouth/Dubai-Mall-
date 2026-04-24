import os
from docx import Document
from docx.shared import Pt, Inches

def create_docx():
    doc = Document()
    
    # Title
    doc.add_heading('Dubai Mall Deck - Source Code', 0)
    
    files_to_add = [
        'index.html',
        'tailwind.config.js',
        'src/index.css',
        'src/App.jsx',
        'src/data/slides.config.js',
        'src/hooks/useCountUp.js',
        'src/utils/icons.jsx',
        'src/components/ChapterPanel.jsx',
        'src/components/SlideNav.jsx',
        'src/components/SlideRenderer.jsx',
        'src/components/SubDeck.jsx',
        'src/slides/ActionClosingSlide.jsx',
        'src/slides/CinematicHeroSlide.jsx',
        'src/slides/EditorialLuxurySlide.jsx',
        'src/slides/GridGallerySlide.jsx',
        'src/slides/MultiColumnSlide.jsx',
        'src/slides/SpecsCardSlide.jsx',
        'src/slides/SplitStatsSlide.jsx'
    ]
    
    for file_path in files_to_add:
        if os.path.exists(file_path):
            # Heading for the file
            heading = doc.add_heading(file_path, level=1)
            
            # Read content
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Add code content
            para = doc.add_paragraph(content)
            para.style.font.name = 'Courier New'
            para.style.font.size = Pt(9)
            
            doc.add_page_break()
    
    output_path = 'Dubai_Mall_Deck_Source_Code.docx'
    doc.save(output_path)
    print(f"Successfully created {output_path}")

if __name__ == '__main__':
    create_docx()
